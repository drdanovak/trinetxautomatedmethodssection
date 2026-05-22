"""
TriNetX Methods Section Generator (Scaffolder) — enhanced version

A wizard-driven Streamlit tool that produces a STROBE/RECORD-aligned first draft
of the Methods section for a TriNetX observational study.

Enhancements in this version:
1. Adds explanatory guidance for each major design choice.
2. Parses optional TriNetX Baseline Patient Characteristics exports to summarize
   propensity-score matching balance.
3. Parses optional TriNetX Measures of Association (MOA) exports to infer effect
   estimates, cohort sizes, event counts, risks, and prior-outcome exclusion notes.
4. Adds target trial emulation, time-zero alignment, missing-data, bias-control,
   and sensitivity-analysis language common in stronger TriNetX methods sections.

The output is a starting scaffold only. Every clinical detail, code list, date,
HCO count, network description, and exported statistic must be verified against
the actual TriNetX query and final manuscript tables.
"""

from __future__ import annotations

import csv
import datetime as _dt
import io
import math
import re
from dataclasses import dataclass
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import streamlit as st

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except Exception:
    pd = None
    PANDAS_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


# ============================================================
# Page setup
# ============================================================
st.set_page_config(
    page_title="Methods Section Generator",
    page_icon="📝",
    layout="wide",
)

st.title("📝 TriNetX Methods Section Generator")
st.caption(
    "A guided, STROBE/RECORD-aligned scaffold for writing the Methods section "
    "of a TriNetX observational study. Optional TriNetX exports can be uploaded "
    "to make the draft more precise."
)

with st.expander("Read this first: what this tool can and cannot do", expanded=False):
    st.markdown(
        """
This tool helps authors produce a structured first draft of a Methods section. It is
especially useful for novice TriNetX users who know what they queried but have not
yet translated that design into journal-ready prose.

It can help you document the data source, study design, eligibility criteria, index
date, exposure/comparator construction, outcome windows, propensity-score matching,
effect estimates, sensitivity analyses, and reporting checks.

It cannot verify that the study is methodologically valid, that a code list is clinically
appropriate, that TriNetX executed the query as intended, or that a causal interpretation
is warranted. The final draft still requires clinical review, methods review, and
verification against the actual TriNetX query.
"""
    )


# ============================================================
# Choice schema
# ============================================================
TRINETX_NETWORKS = [
    "TriNetX Research Network",
    "TriNetX Diamond Network",
    "TriNetX Global Collaborative Network",
    "TriNetX Dataworks – USA",
    "TriNetX US Collaborative Network",
    "TriNetX EMEA Collaborative Network",
    "TriNetX LATAM Collaborative Network",
    "TriNetX APAC Collaborative Network",
    "Other (specify in custom text)",
]

STUDY_DESIGNS = [
    "Retrospective cohort study",
    "Retrospective case-control study",
    "Cross-sectional study",
    "Self-controlled case series",
]

USER_DESIGNS = [
    "New-user (incident) design",
    "Prevalent-user design",
    "Not applicable",
]

COMPARATOR_TYPES = [
    "Active comparator",
    "Non-user / unexposed comparator",
    "Historical comparator",
    "Other (specify in custom text)",
]

TIME_ZERO_STRATEGIES = [
    "Exposure initiation / first qualifying treatment",
    "First qualifying diagnosis",
    "Earliest date all cohort criteria were met",
    "Matched comparator assigned analogous index event",
    "Other (specify)",
]

CENSORING_OPTIONS = [
    "End of outcome window",
    "Last recorded healthcare activity in the network",
    "Occurrence of the outcome",
    "Death",
    "End of available data",
]

EFFECT_ESTIMATES = [
    "Risk ratio (RR) with 95% CI",
    "Odds ratio (OR) with 95% CI",
    "Hazard ratio (HR) with 95% CI from Cox proportional hazards",
    "Risk difference with 95% CI",
    "Kaplan–Meier survival with log-rank test",
    "Competing risk / Aalen–Johansen cumulative incidence",
    "Number needed to treat / harm (NNT/NNH)",
]

EFFECT_ESTIMATES_INLINE = {
    "Risk ratio (RR) with 95% CI": "risk ratios (RRs) with 95% confidence intervals",
    "Odds ratio (OR) with 95% CI": "odds ratios (ORs) with 95% confidence intervals",
    "Hazard ratio (HR) with 95% CI from Cox proportional hazards":
        "hazard ratios (HRs) with 95% confidence intervals from Cox proportional hazards models",
    "Risk difference with 95% CI": "risk differences with 95% confidence intervals",
    "Kaplan–Meier survival with log-rank test":
        "Kaplan–Meier survival estimates with log-rank tests",
    "Competing risk / Aalen–Johansen cumulative incidence":
        "Aalen–Johansen cumulative incidence estimates for competing-risk analyses",
    "Number needed to treat / harm (NNT/NNH)":
        "number needed to treat or harm (NNT/NNH)",
}

MULTIPLE_COMPARISONS_METHODS = [
    "No correction applied (single primary outcome)",
    "No correction applied (exploratory secondary outcomes)",
    "Bonferroni",
    "Holm–Bonferroni",
    "Benjamini–Hochberg FDR",
    "Benjamini–Yekutieli FDR",
]

SENSITIVITY_ANALYSES = [
    "E-value calculation for unmeasured confounding",
    "Varying the lookback window",
    "Varying the outcome window",
    "Alternative matching ratio (e.g., 1:2)",
    "Alternative caliper width",
    "Exclusion of patients with prior outcome events",
    "Lagged outcome start / washout after index",
    "Restricting to patients with ≥1 year of prior healthcare contact",
    "Restricting to patients with known race/ethnicity",
    "Healthcare utilization adjustment",
    "Negative control outcome",
    "Negative control exposure",
    "Competing risk analysis with death as competing event",
    "Subgroup analysis by age",
    "Subgroup analysis by sex",
    "Subgroup analysis by race/ethnicity",
    "Subgroup analysis by disease severity",
]

SENSITIVITY_ANALYSES_INLINE = {
    "E-value calculation for unmeasured confounding":
        "E-value calculation for unmeasured confounding",
    "Varying the lookback window": "variation of the baseline lookback window",
    "Varying the outcome window": "variation of the outcome ascertainment window",
    "Alternative matching ratio (e.g., 1:2)": "an alternative matching ratio",
    "Alternative caliper width": "an alternative propensity-score caliper",
    "Exclusion of patients with prior outcome events":
        "exclusion of patients with any prior occurrence of the outcome",
    "Lagged outcome start / washout after index":
        "a lagged outcome-start window after the index date",
    "Restricting to patients with ≥1 year of prior healthcare contact":
        "restriction to patients with at least 1 year of prior healthcare contact",
    "Restricting to patients with known race/ethnicity":
        "restriction to patients with known race/ethnicity",
    "Healthcare utilization adjustment":
        "adjustment for baseline healthcare utilization",
    "Negative control outcome": "negative control outcome analysis",
    "Negative control exposure": "negative control exposure analysis",
    "Competing risk analysis with death as competing event":
        "competing-risk analysis treating death as a competing event",
    "Subgroup analysis by age": "subgroup analysis by age",
    "Subgroup analysis by sex": "subgroup analysis by sex",
    "Subgroup analysis by race/ethnicity": "subgroup analysis by race/ethnicity",
    "Subgroup analysis by disease severity": "subgroup analysis by disease severity",
}

MISSING_DATA_OPTIONS = [
    "TriNetX retained unknown categories for categorical variables where available",
    "Patients missing continuous variables required for exposure definition were excluded",
    "No imputation was performed by TriNetX",
    "Missingness was not explicitly assessed",
    "Other / custom statement",
]


# ============================================================
# Helper functions
# ============================================================
def _t(s: Optional[str], fallback: str = "") -> str:
    if s is None:
        return fallback
    s = str(s).strip()
    return s if s else fallback


def _bracket(s: Optional[str], fallback: str) -> str:
    val = _t(s)
    return val if val else f"[{fallback}]"


def _list_join(items: Sequence[str]) -> str:
    items = [str(i).strip() for i in items if str(i).strip()]
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} and {items[1]}"
    return ", ".join(items[:-1]) + f", and {items[-1]}"


def _safe_float(x: object) -> Optional[float]:
    if x is None:
        return None
    s = str(x).strip().replace("%", "").replace(",", "")
    if not s:
        return None
    try:
        return float(s)
    except Exception:
        return None


def _format_p(p: Optional[float]) -> str:
    if p is None or math.isnan(p):
        return "[p-value]"
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}".rstrip("0").rstrip(".")


def _format_pct(x: Optional[float], already_percent: bool = False) -> str:
    if x is None or math.isnan(x):
        return "[risk]"
    val = x if already_percent else x * 100
    if abs(val) < 0.01 and val != 0:
        return f"{val:.4f}%"
    if abs(val) < 1:
        return f"{val:.2f}%"
    return f"{val:.1f}%"


def _format_num(x: object) -> str:
    val = _safe_float(x)
    if val is None:
        return str(x)
    if abs(val - int(val)) < 1e-9:
        return f"{int(val):,}"
    return f"{val:,.3f}".rstrip("0").rstrip(".")


def _read_upload_text(uploaded_file) -> str:
    raw = uploaded_file.getvalue()
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return raw.decode(encoding)
        except Exception:
            continue
    return raw.decode("utf-8", errors="replace")


def _csv_rows(text: str) -> List[List[str]]:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    reader = csv.reader(io.StringIO(text))
    rows = []
    for row in reader:
        cleaned = [str(c).strip() for c in row]
        rows.append(cleaned)
    return rows


def _is_blank_row(row: Sequence[str]) -> bool:
    if not row:
        return True
    return all(str(c).strip() in {"", '" "', " "} for c in row)


def _find_section_table(rows: List[List[str]], section_label: str, header_first_cell: str) -> Optional[List[List[str]]]:
    section_label_l = section_label.lower().strip()
    header_first_l = header_first_cell.lower().strip()
    start_idx = None
    for i, row in enumerate(rows):
        if row and row[0].strip().lower() == section_label_l:
            start_idx = i
            break
    if start_idx is None:
        return None

    header_idx = None
    for j in range(start_idx + 1, min(len(rows), start_idx + 12)):
        if rows[j] and rows[j][0].strip().lower() == header_first_l:
            header_idx = j
            break
    if header_idx is None:
        return None

    out = [rows[header_idx]]
    for k in range(header_idx + 1, len(rows)):
        row = rows[k]
        if _is_blank_row(row):
            break
        if len(row) == 1 and row[0].strip().lower() in {
            "risk difference", "risk ratio", "odds ratio", "cohort statistics",
            "baseline patient characteristics"
        }:
            break
        out.append(row)
    return out


def _table_to_dataframe(table_rows: Optional[List[List[str]]]):
    if not table_rows or not PANDAS_AVAILABLE:
        return None
    header = table_rows[0]
    width = len(header)
    body = []
    for row in table_rows[1:]:
        r = list(row[:width])
        if len(r) < width:
            r.extend([""] * (width - len(r)))
        body.append(r)
    try:
        return pd.DataFrame(body, columns=header)
    except Exception:
        return None


def _find_scalar_table(rows: List[List[str]], label: str) -> Optional[Dict[str, str]]:
    label_l = label.lower().strip()
    for i, row in enumerate(rows):
        if row and row[0].strip().lower() == label_l:
            # Find next nonblank row as header and following nonblank row as values.
            header = None
            values = None
            for j in range(i + 1, min(i + 8, len(rows))):
                if not _is_blank_row(rows[j]):
                    header = rows[j]
                    break
            if header is None:
                return None
            for k in range(j + 1, min(j + 8, len(rows))):
                if not _is_blank_row(rows[k]):
                    values = rows[k]
                    break
            if values is None:
                return None
            result = {}
            for idx, h in enumerate(header):
                result[str(h).strip()] = values[idx].strip() if idx < len(values) else ""
            return result
    return None


@dataclass
class PSMSummary:
    filename: str
    n_rows: int = 0
    n_unique_terms: int = 0
    before_max_smd: Optional[float] = None
    after_max_smd: Optional[float] = None
    before_above_threshold: Optional[int] = None
    after_above_threshold: Optional[int] = None
    threshold: float = 0.1
    cohort1_before_n: Optional[int] = None
    cohort2_before_n: Optional[int] = None
    cohort1_after_n: Optional[int] = None
    cohort2_after_n: Optional[int] = None
    top_before: List[Tuple[str, float]] = None
    top_after: List[Tuple[str, float]] = None
    covariate_text: str = ""
    parse_warning: str = ""

    def as_method_text(self) -> str:
        if self.parse_warning:
            return ""
        bits = []
        if self.cohort1_after_n and self.cohort2_after_n:
            bits.append(
                f"After matching, the uploaded baseline-characteristics export contained "
                f"{_format_num(self.cohort1_after_n)} patients in Cohort 1 and "
                f"{_format_num(self.cohort2_after_n)} patients in Cohort 2."
            )
        if self.before_max_smd is not None and self.after_max_smd is not None:
            bits.append(
                f"The maximum absolute standardized mean difference decreased from "
                f"{self.before_max_smd:.3f} before matching to {self.after_max_smd:.3f} "
                f"after matching."
            )
        if self.after_above_threshold is not None:
            if self.after_above_threshold == 0:
                bits.append(
                    f"All measured baseline covariates in the uploaded export were below "
                    f"the prespecified SMD threshold of {self.threshold:g} after matching."
                )
            else:
                bits.append(
                    f"{self.after_above_threshold} measured baseline covariate rows remained "
                    f"above the prespecified SMD threshold of {self.threshold:g} after matching."
                )
        return " ".join(bits)

    def as_qc_markdown(self) -> str:
        if self.parse_warning:
            return f"**{self.filename}:** {self.parse_warning}"
        lines = [
            f"**{self.filename}**",
            f"- Rows parsed: {self.n_rows}",
            f"- Unique covariate terms: {self.n_unique_terms}",
        ]
        if self.cohort1_after_n and self.cohort2_after_n:
            lines.append(
                f"- Matched cohort sizes: Cohort 1 n={_format_num(self.cohort1_after_n)}, "
                f"Cohort 2 n={_format_num(self.cohort2_after_n)}"
            )
        if self.before_max_smd is not None and self.after_max_smd is not None:
            lines.append(
                f"- Max absolute SMD: before={self.before_max_smd:.3f}, "
                f"after={self.after_max_smd:.3f}"
            )
        if self.before_above_threshold is not None and self.after_above_threshold is not None:
            lines.append(
                f"- Rows above SMD {self.threshold:g}: before={self.before_above_threshold}, "
                f"after={self.after_above_threshold}"
            )
        if self.top_after:
            top = ", ".join([f"{name} ({smd:.3f})" for name, smd in self.top_after[:5]])
            lines.append(f"- Largest post-match SMDs: {top}")
        return "\n".join(lines)


@dataclass
class MOASummary:
    filename: str
    cohort1_name: str = ""
    cohort2_name: str = ""
    cohort1_n: Optional[int] = None
    cohort2_n: Optional[int] = None
    cohort1_events: Optional[int] = None
    cohort2_events: Optional[int] = None
    cohort1_risk: Optional[float] = None
    cohort2_risk: Optional[float] = None
    risk_difference: Optional[float] = None
    rd_ci_low: Optional[float] = None
    rd_ci_high: Optional[float] = None
    p_value: Optional[float] = None
    risk_ratio: Optional[float] = None
    rr_ci_low: Optional[float] = None
    rr_ci_high: Optional[float] = None
    odds_ratio: Optional[float] = None
    or_ci_low: Optional[float] = None
    or_ci_high: Optional[float] = None
    prior_exclusion_note: str = ""
    parse_warning: str = ""

    def inferred_effects(self) -> List[str]:
        out = []
        if self.risk_ratio is not None:
            out.append("Risk ratio (RR) with 95% CI")
        if self.odds_ratio is not None:
            out.append("Odds ratio (OR) with 95% CI")
        if self.risk_difference is not None:
            out.append("Risk difference with 95% CI")
        return out

    def as_methods_hint(self) -> str:
        if self.parse_warning:
            return ""
        estimates = []
        if self.risk_difference is not None:
            estimates.append("risk differences")
        if self.risk_ratio is not None:
            estimates.append("risk ratios")
        if self.odds_ratio is not None:
            estimates.append("odds ratios")
        if not estimates:
            return ""
        text = (
            f"The uploaded Measures of Association export reports "
            f"{_list_join(estimates)} based on cohort-level cumulative incidence."
        )
        if self.prior_exclusion_note:
            text += f" The TriNetX export note states: {self.prior_exclusion_note}"
        return text

    def as_qc_markdown(self) -> str:
        if self.parse_warning:
            return f"**{self.filename}:** {self.parse_warning}"
        lines = [f"**{self.filename}**"]
        if self.cohort1_name or self.cohort2_name:
            lines.append(
                f"- Cohorts: {self.cohort1_name or 'Cohort 1'} vs. "
                f"{self.cohort2_name or 'Cohort 2'}"
            )
        if self.cohort1_n and self.cohort2_n:
            lines.append(
                f"- Analytic patients: Cohort 1 n={_format_num(self.cohort1_n)}, "
                f"Cohort 2 n={_format_num(self.cohort2_n)}"
            )
        if self.cohort1_events is not None and self.cohort2_events is not None:
            lines.append(
                f"- Events: Cohort 1 {_format_num(self.cohort1_events)} "
                f"({_format_pct(self.cohort1_risk)}), Cohort 2 "
                f"{_format_num(self.cohort2_events)} ({_format_pct(self.cohort2_risk)})"
            )
        if self.risk_difference is not None:
            lines.append(
                f"- Risk difference: {_format_pct(self.risk_difference)} "
                f"({_format_pct(self.rd_ci_low)} to {_format_pct(self.rd_ci_high)}), "
                f"p={_format_p(self.p_value)}"
            )
        if self.risk_ratio is not None:
            lines.append(
                f"- Risk ratio: {self.risk_ratio:.3f} "
                f"({self.rr_ci_low:.3f} to {self.rr_ci_high:.3f})"
            )
        if self.odds_ratio is not None:
            lines.append(
                f"- Odds ratio: {self.odds_ratio:.3f} "
                f"({self.or_ci_low:.3f} to {self.or_ci_high:.3f})"
            )
        if self.prior_exclusion_note:
            lines.append(f"- Export note: {self.prior_exclusion_note}")
        return "\n".join(lines)


def parse_baseline_characteristics(text: str, filename: str, threshold: float = 0.1) -> PSMSummary:
    summary = PSMSummary(filename=filename, threshold=threshold, top_before=[], top_after=[])
    rows = _csv_rows(text)
    table = _find_section_table(rows, "Baseline Patient Characteristics", "Characteristic ID")
    df = _table_to_dataframe(table)

    if df is None or df.empty:
        summary.parse_warning = "Could not identify a TriNetX Baseline Patient Characteristics table."
        return summary

    summary.n_rows = len(df)
    name_col = "Characteristic Name"
    cat_col = "Category"

    if name_col not in df.columns:
        summary.parse_warning = "The parsed table is missing 'Characteristic Name'."
        return summary

    terms = []
    for _, r in df.iterrows():
        name = str(r.get(name_col, "")).strip()
        cat = str(r.get(cat_col, "")).strip()
        if name and cat:
            terms.append(f"{name}: {cat}")
        elif name:
            terms.append(name)
    summary.n_unique_terms = len(set(terms))
    summary.covariate_text = _list_join(list(dict.fromkeys(terms[:80])))

    before_col = "Before: Standardized Mean Difference"
    after_col = "After: Standardized Mean Difference"
    if before_col in df.columns:
        before_vals = df[before_col].map(_safe_float).dropna().astype(float).abs()
        if not before_vals.empty:
            summary.before_max_smd = float(before_vals.max())
            summary.before_above_threshold = int((before_vals >= threshold).sum())
            top_idx = before_vals.sort_values(ascending=False).head(5).index
            summary.top_before = [(terms[i] if i < len(terms) else f"row {i}", float(before_vals.loc[i])) for i in top_idx]
    if after_col in df.columns:
        after_vals = df[after_col].map(_safe_float).dropna().astype(float).abs()
        if not after_vals.empty:
            summary.after_max_smd = float(after_vals.max())
            summary.after_above_threshold = int((after_vals >= threshold).sum())
            top_idx = after_vals.sort_values(ascending=False).head(5).index
            summary.top_after = [(terms[i] if i < len(terms) else f"row {i}", float(after_vals.loc[i])) for i in top_idx]

    # Use Age at Index row if available to infer cohort sizes; otherwise first row.
    row_for_n = df[df[name_col].astype(str).str.lower().str.contains("age at index", na=False)]
    if row_for_n.empty:
        row_for_n = df.head(1)
    r = row_for_n.iloc[0]
    for attr, col in [
        ("cohort1_before_n", "Cohort 1 Before: Patient Count"),
        ("cohort2_before_n", "Cohort 2 Before: Patient Count"),
        ("cohort1_after_n", "Cohort 1 After: Patient Count"),
        ("cohort2_after_n", "Cohort 2 After: Patient Count"),
    ]:
        if col in df.columns:
            val = _safe_float(r.get(col))
            setattr(summary, attr, int(val) if val is not None else None)

    return summary


def parse_moa_table(text: str, filename: str) -> MOASummary:
    summary = MOASummary(filename=filename)
    rows = _csv_rows(text)

    # Notes section: capture a single useful sentence if present.
    for i, row in enumerate(rows):
        if row and row[0].strip().lower() == "notes":
            notes = []
            for j in range(i + 1, min(i + 8, len(rows))):
                if _is_blank_row(rows[j]):
                    if notes:
                        break
                    continue
                notes.extend([c for c in rows[j] if c.strip()])
            if notes:
                note = " ".join(notes)
                if "excluded" in note.lower():
                    summary.prior_exclusion_note = note
            break

    cohort_table = _find_section_table(rows, "Cohort Statistics", "Cohort")
    df = _table_to_dataframe(cohort_table)
    if df is None or df.empty:
        summary.parse_warning = "Could not identify a TriNetX Cohort Statistics section."
        return summary

    try:
        # Keep only rows with cohort numbers.
        c1 = df[df["Cohort"].astype(str).str.strip() == "1"].iloc[0]
        c2 = df[df["Cohort"].astype(str).str.strip() == "2"].iloc[0]
        summary.cohort1_name = str(c1.get("Cohort Name", "")).strip()
        summary.cohort2_name = str(c2.get("Cohort Name", "")).strip()
        summary.cohort1_n = int(_safe_float(c1.get("Patients in Cohort")) or 0)
        summary.cohort2_n = int(_safe_float(c2.get("Patients in Cohort")) or 0)
        summary.cohort1_events = int(_safe_float(c1.get("Patients with Outcome")) or 0)
        summary.cohort2_events = int(_safe_float(c2.get("Patients with Outcome")) or 0)
        summary.cohort1_risk = _safe_float(c1.get("Risk"))
        summary.cohort2_risk = _safe_float(c2.get("Risk"))
    except Exception:
        summary.parse_warning = "Found Cohort Statistics but could not parse both cohort rows."
        return summary

    rd = _find_scalar_table(rows, "Risk Difference")
    if rd:
        summary.risk_difference = _safe_float(rd.get("Risk Difference"))
        summary.rd_ci_low = _safe_float(rd.get("95 % CI Lower"))
        summary.rd_ci_high = _safe_float(rd.get("95 % CI Upper"))
        summary.p_value = _safe_float(rd.get("p"))

    rr = _find_scalar_table(rows, "Risk Ratio")
    if rr:
        summary.risk_ratio = _safe_float(rr.get("Risk Ratio"))
        summary.rr_ci_low = _safe_float(rr.get("95 % CI Lower"))
        summary.rr_ci_high = _safe_float(rr.get("95 % CI Upper"))

    odds = _find_scalar_table(rows, "Odds Ratio")
    if odds:
        summary.odds_ratio = _safe_float(odds.get("Odds Ratio"))
        summary.or_ci_low = _safe_float(odds.get("95 % CI Lower"))
        summary.or_ci_high = _safe_float(odds.get("95 % CI Upper"))

    return summary


def guide_box(title: str, body: str) -> None:
    with st.expander(title, expanded=False):
        st.markdown(body)


# ============================================================
# Methods paragraph generators
# ============================================================
def para_study_purpose(c: Dict) -> Optional[Dict]:
    if not c.get("include_purpose_statement"):
        return None
    question = _bracket(c.get("research_question"), "state the research question")
    hypotheses = _t(c.get("hypotheses"))
    text = f"This study examined {question}."
    if hypotheses:
        text += f" We prespecified the following hypothesis or analytic expectation: {hypotheses}."
    return {
        "heading": "Study objective and design rationale",
        "text": text,
        "strobe_items": ["2 (Background/rationale)", "3 (Objectives)"],
    }


def para_data_source(c: Dict) -> Dict:
    network = c["network"]
    if network == "Other (specify in custom text)":
        network = _bracket(c.get("network_other"), "specify TriNetX network")

    hco_count = _bracket(c.get("hco_count"), "number")
    patient_pool = _t(c.get("patient_pool"))
    query_date = _bracket(c.get("query_date"), "query date")
    geography = _t(c.get("geography"), "primarily in the United States")
    data_refresh = _t(c.get("data_refresh"))
    record_depth = _t(c.get("record_depth"))
    data_domains = _t(
        c.get("data_domains"),
        "demographics, diagnoses, procedures, medications, laboratory measurements, and encounter information"
    )

    pool_phrase = f", with a total patient pool of approximately {patient_pool}" if patient_pool else ""
    text = (
        f"Data were obtained from the {network}, a federated real-world data network "
        f"providing access to de-identified electronic health record (EHR) data from "
        f"approximately {hco_count} healthcare organizations located {geography}{pool_phrase}. "
        f"Available structured data included {data_domains}. Analyses were conducted in "
        f"the TriNetX Analytics platform on {query_date}."
    )

    if record_depth:
        text += f" The typical longitudinal depth of available records was described as {record_depth}."
    if data_refresh:
        text += f" Data refresh timing was described as {data_refresh}."

    irb_status = c.get("irb_status")
    if irb_status == "IRB exempt / non-human subjects research":
        exempt_basis = _t(
            c.get("irb_exempt_basis"),
            "only aggregated, de-identified results were accessed"
        )
        text += (
            f" The study was determined to be exempt or non-human subjects research because "
            f"{exempt_basis}; informed consent was not required."
        )
    elif irb_status == "IRB approved":
        irb_name = _bracket(c.get("irb_name"), "name of IRB")
        irb_number = _bracket(c.get("irb_number"), "approval number")
        text += f" The study protocol was reviewed and approved by the {irb_name} ({irb_number})."
    else:
        text += " The ethics/IRB status should be verified before manuscript submission."

    reporting_guideline = c.get("reporting_guideline")
    if reporting_guideline:
        text += f" Reporting was guided by the {reporting_guideline} checklist."

    return {
        "heading": "Data source and ethics",
        "text": text,
        "strobe_items": ["4 (Study design)", "5 (Setting)", "22 (Funding/ethics where applicable)"],
    }


def para_design_and_time_zero(c: Dict) -> Dict:
    design = c["study_design"].lower()
    user_design = c.get("user_design", "Not applicable")
    index_event = _bracket(c.get("index_event"), "describe index event")
    time_zero_strategy = c.get("time_zero_strategy", "")
    time_zero_other = _t(c.get("time_zero_other"))
    study_period_start = _bracket(c.get("study_period_start"), "start date")
    study_period_end = _bracket(c.get("study_period_end"), "end date")

    text = (
        f"We conducted a {design}. The study period spanned {study_period_start} to "
        f"{study_period_end}. The index date, or time zero, was defined as {index_event}."
    )

    if time_zero_strategy and time_zero_strategy != "Other (specify)":
        text += f" This corresponds to the following time-zero strategy: {time_zero_strategy.lower()}."
    elif time_zero_other:
        text += f" Time zero was additionally defined as follows: {time_zero_other}."

    if c.get("target_trial_language"):
        text += (
            " The design was informed by target trial emulation principles, including explicit "
            "eligibility criteria, aligned time zero, prespecified exposure strategies, and "
            "prospective follow-up from the index date. Because treatment was not randomized "
            "and not all time-varying confounding could be addressed within the platform, "
            "estimates are interpreted as adjusted associations rather than definitive causal effects."
        )

    if user_design == "New-user (incident) design":
        washout = _bracket(c.get("washout"), "washout window, e.g., 365 days")
        text += (
            f" A new-user design was used: patients were required to have no record of the "
            f"exposure of interest during the {washout} before the index date. This approach "
            f"was selected to reduce prevalent-user bias and improve alignment of exposure "
            f"initiation with the start of follow-up."
        )
    elif user_design == "Prevalent-user design":
        text += (
            " A prevalent-user design was used; patients with prior exposure could enter the "
            "exposed cohort. This choice should be interpreted in light of potential depletion "
            "of susceptibles and survivor bias."
        )

    if c.get("immortal_time_statement"):
        text += (
            " To reduce immortal-time bias, follow-up began only after patients satisfied the "
            "full cohort definition, and the same temporal logic was applied to comparator patients."
        )

    if c.get("study_design_rationale"):
        text += " " + _t(c["study_design_rationale"])

    return {
        "heading": "Study design and time zero",
        "text": text,
        "strobe_items": ["4 (Study design)", "5 (Setting/dates)", "9 (Bias)"],
    }


def para_eligibility(c: Dict) -> Dict:
    age_min = _bracket(c.get("age_min"), "minimum age")
    age_max = _t(c.get("age_max"))
    age_str = f"≥{age_min} years" if not age_max else f"{age_min}–{age_max} years"
    inclusion = _bracket(c.get("inclusion_criteria"), "list inclusion criteria including codes")
    exclusion = _bracket(c.get("exclusion_criteria"), "list exclusion criteria including codes")

    text = (
        f"Eligible patients were aged {age_str} at the index date. Inclusion criteria were: "
        f"{inclusion}. Exclusion criteria were: {exclusion}."
    )

    if c.get("require_prior_encounter"):
        prior_window = _bracket(c.get("prior_encounter_window"), "lookback window")
        text += (
            f" To improve capture of baseline covariates and reduce differential observability, "
            f"patients were required to have at least one healthcare encounter within "
            f"{prior_window} before the index date."
        )

    if c.get("exclude_prior_outcome"):
        text += (
            " Patients with a record of the primary outcome before or on the index date were "
            "excluded to support incident outcome ascertainment."
        )

    return {
        "heading": "Study population and eligibility",
        "text": text,
        "strobe_items": ["6 (Participants)", "7 (Variables)", "8 (Data sources/measurement)"],
    }


def para_exposure_comparator(c: Dict) -> Dict:
    exposure_name = _bracket(c.get("exposure_name"), "exposure name")
    exposure_codes = _bracket(c.get("exposure_codes"), "exposure code list")
    exposure_timing = _t(c.get("exposure_timing"), "on or after the index date")
    comparator_name = _bracket(c.get("comparator_name"), "comparator name")
    comparator_type = c.get("comparator_type", "Active comparator")

    if comparator_type == "Active comparator":
        comp_phrase = f"an active-comparator cohort comprising {comparator_name}"
    elif comparator_type == "Non-user / unexposed comparator":
        comp_phrase = f"an unexposed comparator cohort comprising {comparator_name}"
    elif comparator_type == "Historical comparator":
        comp_phrase = f"a historical comparator cohort comprising {comparator_name}"
    else:
        comp_phrase = _bracket(c.get("comparator_other"), "describe comparator cohort")

    text = (
        f"The exposed cohort consisted of patients with a record of {exposure_name} "
        f"({exposure_codes}) {exposure_timing}. The exposed cohort was compared with "
        f"{comp_phrase}."
    )

    if c.get("active_comparator_rationale") and comparator_type == "Active comparator":
        text += (
            " An active comparator was selected to reduce confounding by indication and "
            "healthcare-seeking behavior relative to a non-user comparator."
        )

    if c.get("require_two_codes"):
        text += (
            " To improve exposure-definition specificity, patients were required to have at "
            "least two records of the qualifying exposure code on separate dates."
        )

    exposure_classification = _t(c.get("exposure_classification"))
    if exposure_classification:
        text += f" Exposure classification followed the following rule: {exposure_classification}."

    return {
        "heading": "Exposure and comparator definitions",
        "text": text,
        "strobe_items": ["7 (Variables)", "8 (Data sources/measurement)", "9 (Bias)"],
    }


def para_outcomes(c: Dict) -> Dict:
    primary_outcome = _bracket(c.get("primary_outcome"), "primary outcome with codes")
    secondary_outcomes = _t(c.get("secondary_outcomes"))
    outcome_window_start = _bracket(c.get("outcome_window_start"), "start of outcome window")
    outcome_window_end = _bracket(c.get("outcome_window_end"), "end of outcome window")
    censoring = c.get("censoring_options", [])

    text = (
        f"The primary outcome was {primary_outcome}. Outcomes were ascertained from "
        f"{outcome_window_start} to {outcome_window_end} relative to the index date."
    )

    if secondary_outcomes:
        text += f" Secondary outcomes included: {secondary_outcomes}."

    if c.get("outcome_first_only"):
        text += " Only the first occurrence of each outcome was counted for each patient."

    if censoring:
        text += f" Follow-up was censored at {_list_join([x.lower() for x in censoring])}."

    if c.get("outcome_validation_statement"):
        text += (
            " Outcome definitions were based on structured diagnosis, procedure, medication, "
            "laboratory, or mortality fields available in TriNetX; code lists should be reported "
            "in the manuscript or supplemental material."
        )

    moa_hints: List[str] = c.get("moa_methods_hints", [])
    if moa_hints and c.get("include_moa_prior_exclusion_note"):
        notes = [h for h in moa_hints if "export note states" in h.lower()]
        if notes:
            note_text = notes[0].split(" The TriNetX export note states: ", 1)[-1]
            text += " The uploaded TriNetX export note states that " + note_text

    return {
        "heading": "Outcome definitions and follow-up",
        "text": text,
        "strobe_items": ["7 (Variables)", "8 (Data sources/measurement)"],
    }


def para_covariates_matching(c: Dict) -> Dict:
    covariates = _bracket(
        c.get("covariates"),
        "list all covariates used for matching or adjustment"
    )
    smd_threshold = _t(c.get("smd_threshold"), "0.1")
    matching_ratio = _t(c.get("matching_ratio"), "1:1")
    caliper = _t(c.get("caliper"), "0.1 standard deviations of the logit of the propensity score")
    baseline_window = _t(c.get("baseline_covariate_window"))
    psm_summary_text = _t(c.get("psm_summary_text"))

    if c.get("use_psm"):
        text = (
            f"To reduce measured confounding and baseline imbalance between cohorts, "
            f"propensity-score matching was performed within the TriNetX Analytics platform. "
            f"Propensity scores were estimated using logistic regression with the following "
            f"baseline covariates: {covariates}."
        )
        if baseline_window:
            text += f" Baseline covariates were assessed during {baseline_window}."
        text += (
            f" Patients were matched {matching_ratio} using greedy nearest-neighbor matching "
            f"with a caliper of {caliper}. Covariate balance was assessed using standardized "
            f"mean differences (SMDs), with absolute SMD values <{smd_threshold} interpreted "
            f"as adequate balance."
        )

        if c.get("missing_data_statement"):
            text += " " + c["missing_data_statement"]

        if psm_summary_text and c.get("include_psm_diagnostics_in_draft"):
            text += " " + psm_summary_text

        if c.get("report_love_plot"):
            text += " Pre- and post-match balance was summarized visually using a Love plot."

    elif c.get("use_adjustment"):
        text = (
            f"To reduce measured confounding, multivariable adjustment was performed using "
            f"the following covariates: {covariates}."
        )
        if baseline_window:
            text += f" Covariates were assessed during {baseline_window}."
        if c.get("missing_data_statement"):
            text += " " + c["missing_data_statement"]
    else:
        text = (
            "Crude, unadjusted comparisons between cohorts were performed. This approach "
            "does not address measured baseline imbalance, so observed associations should "
            "be interpreted descriptively and with caution."
        )

    return {
        "heading": "Covariates and confounding control",
        "text": text,
        "strobe_items": ["7 (Variables)", "9 (Bias)", "12 (Statistical methods)"],
    }


def para_statistical_analysis(c: Dict) -> Dict:
    effects = c.get("effect_estimates", [])
    if not effects:
        effects_str = "[specify which effect estimates were reported]"
    else:
        effects_str = _list_join([EFFECT_ESTIMATES_INLINE.get(e, e) for e in effects])

    alpha = _t(c.get("alpha"), "0.05")
    sided = _t(c.get("sided"), "two-sided")

    text = f"For each outcome, we reported {effects_str}. Statistical significance was defined as a {sided} p-value <{alpha}."

    if any("Risk ratio" in e or "Odds ratio" in e or "Risk difference" in e for e in effects):
        text += (
            " Measures of association for cumulative incidence were estimated using the "
            "TriNetX outcomes analytic."
        )

    if any("Kaplan" in e or "Hazard" in e for e in effects):
        text += (
            " For time-to-event analyses, Kaplan–Meier curves were generated and compared "
            "using log-rank tests; hazard ratios were estimated using Cox proportional hazards "
            "models. The proportional hazards assumption was assessed using available platform "
            "diagnostics and visual inspection of survival curves where appropriate."
        )

    if any("Competing risk" in e or "Aalen" in e for e in effects):
        competing_event = _t(c.get("competing_event"), "death")
        text += (
            f" For competing-risk analyses, {competing_event} was treated as a competing "
            f"event and cumulative incidence was estimated using the Aalen–Johansen estimator."
        )

    mc_method = c.get("multiple_comparisons")
    if mc_method and mc_method.startswith("No correction applied"):
        if "exploratory" in mc_method.lower():
            text += (
                " No formal multiple-comparisons correction was applied to exploratory "
                "secondary outcomes; these analyses were interpreted as hypothesis-generating."
            )
        else:
            text += (
                " A single primary outcome was prespecified, and no multiple-comparisons "
                "correction was applied to the primary analysis."
            )
    elif mc_method:
        family = _t(c.get("multiple_comparisons_family"), "the prespecified family of outcomes")
        text += f" To account for multiple testing across {family}, p-values were adjusted using the {mc_method} method."

    sensitivities = c.get("sensitivity_analyses", [])
    if sensitivities:
        sens_str = _list_join([SENSITIVITY_ANALYSES_INLINE.get(s, s) for s in sensitivities])
        text += f" Prespecified sensitivity and subgroup analyses included {sens_str}."

    custom_stats = _t(c.get("custom_statistical_notes"))
    if custom_stats:
        text += " " + custom_stats

    return {
        "heading": "Statistical analysis",
        "text": text,
        "strobe_items": ["12 (Statistical methods)", "16 (Main results)"],
    }


def para_bias_limitations_in_methods(c: Dict) -> Optional[Dict]:
    if not c.get("include_bias_methods_paragraph"):
        return None
    concerns = c.get("bias_controls", [])
    custom = _t(c.get("custom_bias_text"))
    if not concerns and not custom:
        return None

    concern_text = _list_join([x.lower() for x in concerns])
    text = "Several design choices were used to reduce bias in this observational EHR-based analysis."
    if concern_text:
        text += f" These choices addressed {concern_text}."
    if custom:
        text += " " + custom
    text += (
        " Because residual confounding, outcome misclassification, differential healthcare "
        "utilization, and incomplete capture of care outside contributing organizations may remain, "
        "the estimates should be interpreted as associations."
    )
    return {
        "heading": "Bias-control considerations",
        "text": text,
        "strobe_items": ["9 (Bias)", "19 (Limitations)"],
    }


def para_software(c: Dict) -> Dict:
    toolkit_version = _t(c.get("toolkit_version"), "current version at time of analysis")
    platform_versions = _t(c.get("platform_versions"))

    text = "Analyses were conducted within the TriNetX Analytics web platform."
    if platform_versions:
        text += f" The TriNetX analytic environment was described as follows: {platform_versions}."
    text += (
        f" Manuscript-ready tables, figures, and reporting diagnostics were generated using "
        f"the TriNetX Publication Toolkit ({toolkit_version}), a Streamlit-based application "
        f"for formatting TriNetX exports into publication outputs."
    )
    return {
        "heading": "Software and reporting tools",
        "text": text,
        "strobe_items": ["12 (Statistical methods)", "STROBE/RECORD documentation"],
    }


def build_all_paragraphs(c: Dict) -> List[Dict]:
    candidates = [
        para_study_purpose(c),
        para_data_source(c),
        para_design_and_time_zero(c),
        para_eligibility(c),
        para_exposure_comparator(c),
        para_outcomes(c),
        para_covariates_matching(c),
        para_statistical_analysis(c),
        para_bias_limitations_in_methods(c),
        para_software(c),
    ]
    return [p for p in candidates if p]


# ============================================================
# DOCX export
# ============================================================
def export_docx(paragraphs: List[Dict], title: str = "Methods") -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed.")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.size = Pt(14)

    p_prov = doc.add_paragraph()
    r = p_prov.add_run(
        f"Draft generated by the TriNetX Publication Toolkit Methods Section Generator "
        f"on {_dt.date.today().isoformat()}. Verify every clinical detail, code list, "
        f"date, network descriptor, and statistic before submission."
    )
    r.italic = True
    r.font.size = Pt(9)

    for p in paragraphs:
        doc.add_heading(p["heading"], level=2)
        body = doc.add_paragraph(p["text"])
        body.paragraph_format.space_after = Pt(8)

    doc.add_heading("Verification checklist", level=2)
    checklist_items = [
        "The TriNetX network, HCO count, patient pool, and query date match the final query.",
        "The code lists in the text match the actual cohort definitions and outcome definitions.",
        "The index date, lookback window, outcome window, and censoring rules match the platform query.",
        "The covariates listed are exactly those used for PSM or adjustment.",
        "The PSM balance statements match the final baseline-characteristics export.",
        "The MOA/KM/HR statements match the final TriNetX exports.",
        "The causal language is appropriate for the design and sensitivity analyses performed.",
    ]
    for item in checklist_items:
        doc.add_paragraph(item, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ============================================================
# Optional upload parsing
# ============================================================
st.header("Step 0 · Optional TriNetX export imports")

guide_box(
    "Why upload TriNetX exports?",
    """
The generator works without uploads, but uploaded exports can reduce manual transcription errors.

Upload a **Baseline Patient Characteristics** CSV to extract the covariate list, pre/post-match SMDs,
and matched cohort sizes. Upload one or more **Measures of Association** CSVs to infer whether the
analysis used risk difference, risk ratio, odds ratio, event counts, cohort sizes, and prior-outcome
exclusions.

The imports do not prove that the design is valid. They only help the wizard draft language that is
consistent with the files you provide.
"""
)

psm_upload = st.file_uploader(
    "Optional: upload TriNetX Baseline Patient Characteristics CSV",
    type=["csv", "txt"],
    help="Use the CSV exported from TriNetX after propensity score matching. The tool will parse pre/post SMDs and matched cohort sizes.",
)

moa_uploads = st.file_uploader(
    "Optional: upload one or more TriNetX Measures of Association CSVs",
    type=["csv", "txt"],
    accept_multiple_files=True,
    help="Use the MOA table export for one or more outcomes. The tool will parse cohort sizes, event counts, risk, RR, OR, RD, CIs, and p-values.",
)

parsed_psm: Optional[PSMSummary] = None
parsed_moas: List[MOASummary] = []
auto_effects: List[str] = []

# The SMD threshold control appears later; parse with conventional default here.
if psm_upload is not None:
    parsed_psm = parse_baseline_characteristics(_read_upload_text(psm_upload), psm_upload.name, threshold=0.1)
    st.markdown(parsed_psm.as_qc_markdown())

if moa_uploads:
    for f in moa_uploads:
        parsed = parse_moa_table(_read_upload_text(f), f.name)
        parsed_moas.append(parsed)
        auto_effects.extend(parsed.inferred_effects())
    with st.expander("Parsed Measures of Association summaries", expanded=True):
        for parsed in parsed_moas:
            st.markdown(parsed.as_qc_markdown())
            st.divider()

auto_effects = list(dict.fromkeys(auto_effects))


# ============================================================
# UI: Step 1
# ============================================================
st.divider()
st.header("Step 1 · Study purpose and design intent")

guide_box(
    "How to answer this section",
    """
Strong TriNetX methods sections usually begin with a clear research question and then describe the
design choices that make the EHR analysis interpretable. Use this section to state the question, the
main exposure/comparator contrast, and whether the design is meant to approximate a target trial.

Use target-trial language when you have deliberately aligned eligibility, treatment assignment, time
zero, follow-up, and outcomes. Do not use causal language if the design is primarily descriptive.
"""
)

include_purpose_statement = st.checkbox(
    "Include a brief study objective/hypothesis paragraph",
    value=True,
    help="Recommended for manuscripts where the Methods section begins by restating the analytic question or prespecified hypotheses.",
)
research_question = ""
hypotheses = ""
if include_purpose_statement:
    research_question = st.text_area(
        "Research question / objective",
        placeholder="e.g., whether GLP-1 receptor agonist initiation is associated with incident autoimmune disease among adults with type 2 diabetes",
        height=70,
        help="Write this as a methods-ready object of analysis, not a broad background rationale.",
    )
    hypotheses = st.text_area(
        "Prespecified hypothesis or analytic expectation (optional)",
        placeholder="e.g., We hypothesized that GLP-1 RA initiators would have different rates of incident autoimmune outcomes than active-comparator initiators.",
        height=70,
        help="Include only if the analysis was actually prespecified. Leave blank for exploratory studies.",
    )

target_trial_language = st.checkbox(
    "Include target trial emulation language",
    value=True,
    help="Use when the design explicitly aligns eligibility criteria, treatment strategies, time zero, and follow-up. The generated text will still avoid overclaiming causality.",
)


# ============================================================
# UI: Step 2
# ============================================================
st.divider()
st.header("Step 2 · Data source and approvals")

guide_box(
    "What reviewers look for in the data source paragraph",
    """
Reviewers expect the methods to identify the TriNetX network, approximate HCO count, patient pool
when relevant, geography, query date, available structured data domains, and ethics status. Strong
papers also state that participating organizations are de-identified and that the analysis used
structured EHR fields, not manual chart review.
"""
)

col1, col2 = st.columns(2)
with col1:
    network = st.selectbox(
        "TriNetX network used",
        TRINETX_NETWORKS,
        index=4 if "TriNetX US Collaborative Network" in TRINETX_NETWORKS else 0,
        help="Use the exact network named in your TriNetX query. Network names change; verify before submission.",
    )
    network_other = ""
    if network == "Other (specify in custom text)":
        network_other = st.text_input("Specify network name")
    hco_count = st.text_input(
        "Approximate number of contributing HCOs at query time",
        placeholder="e.g., 72",
        help="Use the number visible in TriNetX or institutional documentation on the final query date.",
    )
    patient_pool = st.text_input(
        "Approximate patient pool at query time (optional)",
        placeholder="e.g., 126,379,149",
        help="Often useful for describing the network, especially in high-impact methods sections.",
    )
    geography = st.text_input(
        "Geographic coverage",
        placeholder="e.g., in the United States",
        help="Be specific: U.S., global, EMEA, APAC, etc.",
    )
with col2:
    query_date = st.text_input(
        "Date the TriNetX query was executed",
        placeholder="e.g., 4 July 2025",
        help="Use the final date of data collection/query execution.",
    )
    reporting_guideline = st.selectbox(
        "Reporting guideline referenced",
        ["STROBE", "RECORD", "STROBE and RECORD", "None specified"],
        index=2,
        help="RECORD is especially relevant for routinely collected health data; STROBE is widely recognized for observational studies.",
    )
    if reporting_guideline == "None specified":
        reporting_guideline = ""
    irb_status = st.radio(
        "Ethics / IRB status",
        [
            "IRB exempt / non-human subjects research",
            "IRB approved",
            "Not yet determined",
        ],
        index=0,
        help="Use your actual institutional determination. Do not assume exemption without documentation.",
    )
    irb_name = ""
    irb_number = ""
    irb_exempt_basis = ""
    if irb_status == "IRB approved":
        irb_name = st.text_input("IRB name")
        irb_number = st.text_input("IRB approval number")
    elif irb_status == "IRB exempt / non-human subjects research":
        irb_exempt_basis = st.text_input(
            "Basis for exemption / non-human subjects determination",
            value="only aggregated, de-identified results were accessed through TriNetX",
        )

data_domains = st.text_area(
    "Structured data domains available / used",
    value="demographics, diagnoses, procedures, medications, laboratory measurements, encounter information, and mortality information where available",
    height=70,
    help="Mention only domains relevant to the study and available in your network.",
)
col1, col2 = st.columns(2)
with col1:
    record_depth = st.text_input(
        "Longitudinal record depth or censoring description (optional)",
        placeholder="e.g., average record length approximately seven years; records censored at 20 years or last patient fact",
    )
with col2:
    data_refresh = st.text_input(
        "Data refresh statement (optional)",
        placeholder="e.g., data are refreshed periodically by participating HCOs",
    )


# ============================================================
# UI: Step 3
# ============================================================
st.divider()
st.header("Step 3 · Study design, index date, and time zero")

guide_box(
    "How to choose the study design and time-zero options",
    """
Use a **retrospective cohort** design when patients are classified by exposure and followed forward
in the record for incident outcomes. Use a **case-control** design only when cases are selected by
outcome status and compared with controls. Use a **new-user design** when you require no prior
exposure and anchor follow-up at initiation; this usually improves interpretability for treatment
studies. The index date should be the first moment a patient fully qualifies for the analytic cohort.
"""
)

col1, col2 = st.columns(2)
with col1:
    study_design = st.selectbox("Study design", STUDY_DESIGNS, index=0)
    user_design = st.selectbox(
        "Exposure-user design",
        USER_DESIGNS,
        index=0,
        help="New-user designs reduce prevalent-user bias when studying treatments; prevalent-user designs may be acceptable for descriptive exposure histories.",
    )
    washout = ""
    if user_design == "New-user (incident) design":
        washout = st.text_input(
            "Exposure washout period before index date",
            placeholder="e.g., 365 days",
            help="State the period during which patients could not have prior exposure.",
        )
with col2:
    study_period_start = st.text_input("Study period start", placeholder="e.g., 1 January 2015")
    study_period_end = st.text_input("Study period end", placeholder="e.g., 31 December 2024")
    time_zero_strategy = st.selectbox(
        "Time-zero strategy",
        TIME_ZERO_STRATEGIES,
        index=2,
        help="Time zero should align exposure/comparator eligibility and start of follow-up.",
    )
    time_zero_other = ""
    if time_zero_strategy == "Other (specify)":
        time_zero_other = st.text_input("Describe time-zero strategy")

index_event = st.text_area(
    "Index event / time-zero definition",
    placeholder="e.g., the earliest date on which both a qualifying dyslipidemia diagnosis and statin prescription were recorded",
    height=70,
    help="This is one of the most important fields. Define exactly when follow-up starts.",
)

immortal_time_statement = st.checkbox(
    "Include explicit immortal-time-bias reduction language",
    value=True,
    help="Recommended when exposure timing could otherwise create immortal person-time.",
)

study_design_rationale = st.text_area(
    "Additional design rationale (optional)",
    placeholder="e.g., We selected an active comparator to reduce channeling bias and aligned the comparator index date to the first qualifying alternative therapy prescription.",
    height=80,
)


# ============================================================
# UI: Step 4
# ============================================================
st.divider()
st.header("Step 4 · Eligibility criteria")

guide_box(
    "How to write eligibility criteria",
    """
Eligibility criteria should be specific enough that another TriNetX user could reproduce the query.
Include age, diagnoses, medications, labs, required encounters, lookback windows, prior outcome
exclusions, and any disease-specific exclusions. Avoid hiding the code list in vague prose; either
include codes here or state that the complete code list appears in a supplement.
"""
)

col1, col2, col3 = st.columns(3)
with col1:
    age_min = st.text_input("Minimum age at index", value="18")
with col2:
    age_max = st.text_input("Maximum age at index (optional)")
with col3:
    require_prior_encounter = st.checkbox(
        "Require ≥1 prior healthcare encounter",
        value=True,
        help="This helps reduce differential missingness and improves baseline covariate capture.",
    )

prior_encounter_window = ""
if require_prior_encounter:
    prior_encounter_window = st.text_input(
        "Prior encounter lookback window",
        value="365 days",
        help="Use the actual lookback window from the query.",
    )

exclude_prior_outcome = st.checkbox(
    "Exclude patients with prior occurrence of the primary outcome",
    value=True,
    help="Recommended for incident-outcome analyses.",
)

inclusion_criteria = st.text_area(
    "Inclusion criteria",
    placeholder="e.g., adults with ≥1 ICD-10-CM E78 diagnosis during the study period and at least one qualifying medication record...",
    height=100,
)
exclusion_criteria = st.text_area(
    "Exclusion criteria",
    placeholder="e.g., prior diagnosis of the primary outcome before or on index; pregnancy-related codes; prior malignancy; comparator medication exposure...",
    height=100,
)


# ============================================================
# UI: Step 5
# ============================================================
st.divider()
st.header("Step 5 · Exposure and comparator")

guide_box(
    "Comparator choice guidance",
    """
An **active comparator** is usually stronger for pharmacoepidemiology because both groups have an
indication for treatment and comparable healthcare contact. A **non-user comparator** can be useful
when no clinically reasonable active comparator exists, but it increases concern for confounding by
indication, healthy-user bias, and differential surveillance. A **historical comparator** should be used
cautiously because secular changes in coding, care, and outcome detection can bias results.
"""
)

col1, col2 = st.columns(2)
with col1:
    exposure_name = st.text_input(
        "Exposure name",
        placeholder="e.g., GLP-1 receptor agonist",
        help="Use the clinical name, not only the code category.",
    )
    exposure_codes = st.text_area(
        "Exposure code list",
        placeholder="e.g., RxNorm codes for semaglutide, liraglutide, dulaglutide...",
        height=80,
    )
    exposure_timing = st.text_input(
        "Exposure timing rule",
        value="on or after the index date",
        help="For new-user designs, specify first qualifying exposure and any required temporal relation to diagnosis.",
    )
    require_two_codes = st.checkbox(
        "Require ≥2 records of the qualifying exposure code on separate dates",
        value=False,
    )
with col2:
    comparator_type = st.selectbox("Comparator type", COMPARATOR_TYPES, index=0)
    comparator_name = st.text_input(
        "Comparator cohort label",
        placeholder="e.g., DPP-4 inhibitor initiators",
    )
    comparator_other = ""
    if comparator_type == "Other (specify in custom text)":
        comparator_other = st.text_input("Describe the comparator")
    active_comparator_rationale = st.checkbox(
        "Include active-comparator rationale",
        value=(comparator_type == "Active comparator"),
    )

exposure_classification = st.text_area(
    "Additional exposure classification rule (optional)",
    placeholder="e.g., Exposure was ascertained once at index and analyzed using an intention-to-treat framework; switching/discontinuation was not time-updated.",
    height=70,
)


# ============================================================
# UI: Step 6
# ============================================================
st.divider()
st.header("Step 6 · Outcomes and follow-up")

guide_box(
    "Outcome-window guidance",
    """
Define outcomes using codes and time windows. For incident diagnoses, consider excluding outcomes
on or before index and starting follow-up after a short lag when reverse causality is plausible. For
mortality or safety outcomes, be explicit about whether events on index day count. If using multiple
outcomes, distinguish prespecified primary outcomes from exploratory secondary outcomes.
"""
)

primary_outcome = st.text_area(
    "Primary outcome",
    placeholder="e.g., incident Alzheimer’s disease defined by ICD-10-CM G30 after the index date",
    height=80,
)

secondary_outcomes = st.text_area(
    "Secondary outcomes (optional)",
    placeholder="e.g., early-onset AD (G30.0), late-onset AD (G30.1), all-cause mortality...",
    height=80,
)

col1, col2, col3 = st.columns(3)
with col1:
    outcome_window_start = st.text_input("Outcome window start", value="1 day after index")
with col2:
    outcome_window_end = st.text_input("Outcome window end", value="5 years after index")
with col3:
    outcome_first_only = st.checkbox("Count first occurrence only", value=True)

censoring_options = st.multiselect(
    "Censoring rules",
    CENSORING_OPTIONS,
    default=["End of outcome window", "Last recorded healthcare activity in the network", "Occurrence of the outcome"],
    help="Select all censoring rules that apply to the analysis.",
)

outcome_validation_statement = st.checkbox(
    "Include statement that structured code lists should be reported in manuscript/supplement",
    value=True,
)

include_moa_prior_exclusion_note = st.checkbox(
    "Use uploaded MOA prior-outcome exclusion note in outcome paragraph when available",
    value=True if parsed_moas else False,
    disabled=False if parsed_moas else True,
)


# ============================================================
# UI: Step 7
# ============================================================
st.divider()
st.header("Step 7 · Confounding control and matching")

guide_box(
    "PSM reporting guidance",
    """
A strong PSM paragraph reports why matching was used, the covariates included, the baseline window
for covariates, the algorithm, matching ratio, caliper, whether replacement was allowed if known, and
how balance was judged. It should also report pre/post-match SMDs somewhere in the manuscript,
usually Table 1, a Love plot, or the Results.
"""
)

confounding_strategy = st.radio(
    "Confounding control strategy",
    ["Propensity score matching", "Multivariable adjustment", "Crude (no adjustment)"],
    index=0,
)
use_psm = confounding_strategy == "Propensity score matching"
use_adjustment = confounding_strategy == "Multivariable adjustment"

auto_covariates = parsed_psm.covariate_text if parsed_psm and parsed_psm.covariate_text else ""
covariates = st.text_area(
    "Covariates used for matching or adjustment",
    value=auto_covariates,
    placeholder="e.g., age at index, sex, race/ethnicity, BMI, smoking, comorbidities, baseline medications, healthcare utilization...",
    height=140,
    help="If a PSM table was uploaded, this field is prefilled from the parsed baseline-characteristics rows. Review and simplify before using in a manuscript.",
)

baseline_covariate_window = st.text_input(
    "Baseline covariate assessment window (optional)",
    placeholder="e.g., 365 days before index; any time before index; on date of index",
)

matching_ratio = ""
caliper = ""
smd_threshold = "0.1"
report_love_plot = False
include_psm_diagnostics_in_draft = False

if use_psm:
    col1, col2, col3 = st.columns(3)
    with col1:
        matching_ratio = st.text_input("Matching ratio", value="1:1")
    with col2:
        caliper = st.text_input(
            "Caliper width",
            value="0.1 standard deviations of the logit of the propensity score",
        )
    with col3:
        smd_threshold = st.text_input("SMD balance threshold", value="0.1")

    report_love_plot = st.checkbox("Report/display Love plot or balance plot", value=True)
    include_psm_diagnostics_in_draft = st.checkbox(
        "Include uploaded PSM balance diagnostic summary in generated draft",
        value=True if parsed_psm else False,
        disabled=False if parsed_psm else True,
        help="Some journals prefer detailed balance results in Results rather than Methods. Use only if this fits your manuscript style.",
    )

missing_options = st.multiselect(
    "Missing-data statements to include",
    MISSING_DATA_OPTIONS,
    default=[
        "No imputation was performed by TriNetX",
    ],
    help="Use only statements that are true for your TriNetX analysis and network.",
)
missing_custom = ""
if "Other / custom statement" in missing_options:
    missing_custom = st.text_input("Custom missing-data statement")

missing_data_parts = [x for x in missing_options if x != "Other / custom statement"]
if missing_custom:
    missing_data_parts.append(missing_custom)
missing_data_statement = ""
if missing_data_parts:
    missing_data_statement = _list_join(missing_data_parts) + "."


# Reparse PSM summary if user changed threshold
psm_summary_text = ""
if parsed_psm and not parsed_psm.parse_warning:
    thr_val = _safe_float(smd_threshold) or 0.1
    # Reparse at selected threshold from uploaded file.
    try:
        parsed_psm = parse_baseline_characteristics(_read_upload_text(psm_upload), psm_upload.name, threshold=thr_val)
    except Exception:
        pass
    psm_summary_text = parsed_psm.as_method_text()


# ============================================================
# UI: Step 8
# ============================================================
st.divider()
st.header("Step 8 · Statistical analysis")

guide_box(
    "Statistical analysis guidance",
    """
TriNetX MOA tables usually support cumulative-incidence statements with risk differences, risk
ratios, odds ratios, confidence intervals, z statistics, and p-values. Time-to-event analyses should
state Kaplan–Meier estimation, log-rank tests, Cox proportional hazards models, follow-up start,
censoring, and the proportional hazards assumption. Multiple-outcome studies should define the
family of tests for any correction.
"""
)

default_effects = auto_effects if auto_effects else [
    "Risk ratio (RR) with 95% CI",
    "Hazard ratio (HR) with 95% CI from Cox proportional hazards",
    "Kaplan–Meier survival with log-rank test",
]
effect_estimates = st.multiselect(
    "Effect estimates reported",
    EFFECT_ESTIMATES,
    default=[e for e in default_effects if e in EFFECT_ESTIMATES],
)

competing_event = ""
if "Competing risk / Aalen–Johansen cumulative incidence" in effect_estimates:
    competing_event = st.text_input("Competing event", value="all-cause mortality")

col1, col2, col3 = st.columns(3)
with col1:
    alpha = st.text_input("Significance threshold (alpha)", value="0.05")
with col2:
    sided = st.selectbox("Test sidedness", ["two-sided", "one-sided"], index=0)
with col3:
    multiple_comparisons = st.selectbox(
        "Multiple comparisons correction",
        MULTIPLE_COMPARISONS_METHODS,
        index=0,
    )

multiple_comparisons_family = ""
if not multiple_comparisons.startswith("No correction"):
    multiple_comparisons_family = st.text_input(
        "Family of tests corrected",
        placeholder="e.g., five prespecified outcomes within each analysis",
    )

sensitivity_analyses = st.multiselect(
    "Prespecified sensitivity / subgroup analyses",
    SENSITIVITY_ANALYSES,
    default=[],
)

custom_statistical_notes = st.text_area(
    "Additional statistical methods notes (optional)",
    placeholder="e.g., Follow-up was repeated at 1-year, 3-year, and 5-year horizons; analyses were repeated in a restricted cohort...",
    height=80,
)


# ============================================================
# UI: Step 9
# ============================================================
st.divider()
st.header("Step 9 · Bias-control language")

guide_box(
    "When to include a bias-control paragraph",
    """
Many stronger TriNetX methods sections explicitly describe how the design addressed known threats:
confounding by indication, reverse causality, immortal-time bias, surveillance bias, missingness,
selection into the network, and outcome misclassification. This paragraph should not claim that bias
was eliminated; it should state what was done and what residual limitations remain.
"""
)

include_bias_methods_paragraph = st.checkbox(
    "Include a short bias-control considerations paragraph",
    value=True,
)
bias_controls = []
custom_bias_text = ""
if include_bias_methods_paragraph:
    bias_controls = st.multiselect(
        "Biases addressed by design",
        [
            "confounding by indication",
            "baseline imbalance",
            "immortal-time bias",
            "reverse causality",
            "surveillance/detection bias",
            "outcome misclassification",
            "selection bias due to incomplete EHR capture",
            "healthy-user bias",
            "depletion of susceptibles",
        ],
        default=[
            "confounding by indication",
            "baseline imbalance",
            "immortal-time bias",
            "reverse causality",
            "surveillance/detection bias",
        ],
    )
    custom_bias_text = st.text_area(
        "Custom bias-control text (optional)",
        placeholder="e.g., We required prior healthcare contact and adjusted for utilization proxies to reduce differential observability between cohorts.",
        height=70,
    )


# ============================================================
# UI: Step 10
# ============================================================
st.divider()
st.header("Step 10 · Software and toolkit")

platform_versions = st.text_area(
    "TriNetX platform analytic environment / version statement (optional)",
    placeholder="e.g., Java 11.0.16; R 4.0.2; Python 3.7; lifelines 0.22.4...",
    height=80,
)
toolkit_version = st.text_input(
    "TriNetX Publication Toolkit version / date",
    value="current version at time of manuscript preparation",
)


# ============================================================
# Consolidate choices
# ============================================================
moa_methods_hints = [m.as_methods_hint() for m in parsed_moas if m.as_methods_hint()]

choices = {
    "include_purpose_statement": include_purpose_statement,
    "research_question": research_question,
    "hypotheses": hypotheses,
    "target_trial_language": target_trial_language,

    "network": network,
    "network_other": network_other,
    "hco_count": hco_count,
    "patient_pool": patient_pool,
    "geography": geography,
    "query_date": query_date,
    "reporting_guideline": reporting_guideline,
    "irb_status": irb_status,
    "irb_name": irb_name,
    "irb_number": irb_number,
    "irb_exempt_basis": irb_exempt_basis,
    "data_domains": data_domains,
    "record_depth": record_depth,
    "data_refresh": data_refresh,

    "study_design": study_design,
    "user_design": user_design,
    "washout": washout,
    "study_period_start": study_period_start,
    "study_period_end": study_period_end,
    "time_zero_strategy": time_zero_strategy,
    "time_zero_other": time_zero_other,
    "index_event": index_event,
    "immortal_time_statement": immortal_time_statement,
    "study_design_rationale": study_design_rationale,

    "age_min": age_min,
    "age_max": age_max,
    "require_prior_encounter": require_prior_encounter,
    "prior_encounter_window": prior_encounter_window,
    "exclude_prior_outcome": exclude_prior_outcome,
    "inclusion_criteria": inclusion_criteria,
    "exclusion_criteria": exclusion_criteria,

    "exposure_name": exposure_name,
    "exposure_codes": exposure_codes,
    "exposure_timing": exposure_timing,
    "require_two_codes": require_two_codes,
    "comparator_type": comparator_type,
    "comparator_name": comparator_name,
    "comparator_other": comparator_other,
    "active_comparator_rationale": active_comparator_rationale,
    "exposure_classification": exposure_classification,

    "primary_outcome": primary_outcome,
    "secondary_outcomes": secondary_outcomes,
    "outcome_window_start": outcome_window_start,
    "outcome_window_end": outcome_window_end,
    "outcome_first_only": outcome_first_only,
    "censoring_options": censoring_options,
    "outcome_validation_statement": outcome_validation_statement,
    "moa_methods_hints": moa_methods_hints,
    "include_moa_prior_exclusion_note": include_moa_prior_exclusion_note,

    "use_psm": use_psm,
    "use_adjustment": use_adjustment,
    "covariates": covariates,
    "baseline_covariate_window": baseline_covariate_window,
    "matching_ratio": matching_ratio,
    "caliper": caliper,
    "smd_threshold": smd_threshold,
    "report_love_plot": report_love_plot,
    "include_psm_diagnostics_in_draft": include_psm_diagnostics_in_draft,
    "missing_data_statement": missing_data_statement,
    "psm_summary_text": psm_summary_text,

    "effect_estimates": effect_estimates,
    "competing_event": competing_event,
    "alpha": alpha,
    "sided": sided,
    "multiple_comparisons": multiple_comparisons,
    "multiple_comparisons_family": multiple_comparisons_family,
    "sensitivity_analyses": sensitivity_analyses,
    "custom_statistical_notes": custom_statistical_notes,

    "include_bias_methods_paragraph": include_bias_methods_paragraph,
    "bias_controls": bias_controls,
    "custom_bias_text": custom_bias_text,

    "platform_versions": platform_versions,
    "toolkit_version": toolkit_version,
}


# ============================================================
# Output
# ============================================================
st.divider()
st.header("Generated Methods draft")

paragraphs = build_all_paragraphs(choices)

for p in paragraphs:
    st.subheader(p["heading"])
    st.markdown(p["text"])
    st.caption("Reporting items addressed: " + ", ".join(p["strobe_items"]))

with st.expander("Plain-text version for copy/paste", expanded=False):
    full_text = ""
    for p in paragraphs:
        full_text += f"\n## {p['heading']}\n\n{p['text']}\n"
    st.code(full_text.strip(), language="markdown")

with st.expander("Export-derived quality-control summaries", expanded=False):
    if parsed_psm:
        st.markdown(parsed_psm.as_qc_markdown())
    if parsed_moas:
        for parsed in parsed_moas:
            st.markdown(parsed.as_qc_markdown())
            st.divider()
    if not parsed_psm and not parsed_moas:
        st.info("No TriNetX exports were uploaded.")


# ============================================================
# Verification gate
# ============================================================
st.divider()
st.header("Verification gate before Word export")

st.markdown(
    "Confirm these items against the actual TriNetX query and exported tables. The generator cannot verify them automatically."
)

v1 = st.checkbox("The TriNetX network, HCO count, patient pool, and query date are correct.")
v2 = st.checkbox("The inclusion, exclusion, exposure, comparator, and outcome code lists match the TriNetX query.")
v3 = st.checkbox("The index date, time-zero logic, lookback window, outcome window, and censoring rules are correct.")
v4 = st.checkbox("The covariates listed are exactly those used for PSM or adjustment.")
v5 = st.checkbox("Any PSM balance and MOA statements match the final TriNetX exports.")
v6 = st.checkbox("The causal language is appropriately cautious for an observational EHR study.")

gate_open = all([v1, v2, v3, v4, v5, v6])

if gate_open and DOCX_AVAILABLE:
    docx_bytes = export_docx(paragraphs)
    st.download_button(
        "Download Methods section as Word document",
        data=docx_bytes,
        file_name="trinetx_methods_draft.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
elif not DOCX_AVAILABLE:
    st.warning("python-docx is not installed; Word export is unavailable.")
else:
    st.info("Tick all verification items above to unlock the Word export.")
